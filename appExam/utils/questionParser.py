import logging
import re
from io import TextIOWrapper
from typing import Any

import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)


class DocumentFormatConfig:
    """Configuration for document parsing format"""

    def __init__(self):
        # CSV format configuration
        self.csv_question_column = "QUESTION"
        self.csv_answer_column = "ANSWER"
        self.csv_option_columns = ["OPTIONS_A", "OPTIONS_B", "OPTIONS_C", "OPTIONS_D"]
        self.csv_answer_format = "letter"  # "letter" or "text"

        # Text format configuration
        self.text_question_prefix = "Q."
        self.text_option_prefixes = ["1.", "2.", "3.", "4."]
        self.text_answer_prefix = "Answer:"
        self.text_answer_format = "text"  # "text" or "number"

        # DOCX format configuration (same as text for now)
        self.docx_question_prefix = "Q."
        self.docx_option_prefixes = ["1.", "2.", "3.", "4."]
        self.docx_answer_prefix = "Answer:"
        self.docx_answer_format = "text"  # "text" or "number"


def parse_questions_from_csv(
    file,
    config: DocumentFormatConfig = None,
) -> list[dict[str, Any]]:
    """
    Parses questions from a CSV file with configurable format
    """
    if config is None:
        config = DocumentFormatConfig()

    try:
        df = pd.read_csv(TextIOWrapper(file, encoding="utf-8"))
        logger.info("CSV file successfully read with %d rows.", len(df))
    except Exception as e:
        logger.exception("Failed to read CSV file: %s", str(e))
        raise

    questions = []
    for idx, row in df.iterrows():
        question_text = str(row.get(config.csv_question_column, "")).strip()
        answer_value = str(row.get(config.csv_answer_column, "")).strip()

        # Get options
        options = {}
        option_letters = ["a", "b", "c", "d"]
        for i, col in enumerate(config.csv_option_columns):
            if i < len(option_letters):
                options[option_letters[i]] = str(row.get(col, "")).strip()

        # Process answer based on format
        if config.csv_answer_format == "letter":
            answer_letter = answer_value.lower()
            if answer_letter not in options:
                logger.warning(
                    "Row %d: Invalid answer '%s'. Must be one of %s.",
                    idx + 1,
                    answer_letter,
                    "/".join(options.keys()),
                )
                continue
        else:  # text format
            answer_letter = None
            for letter, text in options.items():
                if text.lower() == answer_value.lower():
                    answer_letter = letter
                    break

            if answer_letter is None:
                logger.warning(
                    "Row %d: Could not match answer text '%s' to any option.",
                    idx + 1,
                    answer_value,
                )
                continue

        # Build answers list
        answers = []
        for letter, text in options.items():
            if not text:
                continue
            answers.append(
                {
                    "text": text,
                    "option_letter": letter.upper(),
                    "is_correct": (letter == answer_letter),
                },
            )

        if not question_text:
            logger.warning("Row %d: Empty question text.", idx + 1)
        if not answers:
            logger.warning("Row %d: No valid options provided.", idx + 1)

        if question_text and answers:
            questions.append(
                {
                    "question": question_text,
                    "answers": answers,
                },
            )
            logger.debug(
                "Row %d: Parsed question with %d options.",
                idx + 1,
                len(answers),
            )

    logger.info("Total valid questions parsed: %d", len(questions))
    return questions


def parse_questions_from_text(
    file,
    config: DocumentFormatConfig = None,
) -> list[dict[str, Any]]:
    """
    Parses questions from a text file with configurable format
    """
    if config is None:
        config = DocumentFormatConfig()

    try:
        content = file.read().decode("utf-8")
        logger.info("Text file successfully read.")
    except Exception as e:
        logger.exception("Failed to read text file: %s", str(e))
        raise

    # Check if it's the standard MCQ format (Question X) ... Ans.)
    if "Question" in content and "Ans." in content and "A)" in content:
        return _parse_mcq_format(content)

    # Standard configurable format parsing
    questions = []
    current_question = None
    current_options = []
    current_answer = None

    lines = content.split("\n")
    for line_num, line in enumerate(lines, 1):
        line = line.strip()
        if not line:
            continue

        # Check if it's a question
        if line.startswith(config.text_question_prefix):
            # Save previous question if exists
            if current_question and current_options:
                question_data = _build_question_from_text(
                    current_question,
                    current_options,
                    current_answer,
                    config,
                )
                if question_data:
                    questions.append(question_data)

            # Start new question
            current_question = line[len(config.text_question_prefix) :].strip()
            current_options = []
            current_answer = None

        # Check if it's an option
        elif any(line.startswith(prefix) for prefix in config.text_option_prefixes):
            for i, prefix in enumerate(config.text_option_prefixes):
                if line.startswith(prefix):
                    option_text = line[len(prefix) :].strip()
                    current_options.append(
                        {"index": i, "text": option_text, "letter": chr(ord("A") + i)},
                    )
                    break

        # Check if it's an answer
        elif line.startswith(config.text_answer_prefix):
            current_answer = line[len(config.text_answer_prefix) :].strip()

    # Don't forget the last question
    if current_question and current_options:
        question_data = _build_question_from_text(
            current_question,
            current_options,
            current_answer,
            config,
        )
        if question_data:
            questions.append(question_data)

    logger.info("Total valid questions parsed: %d", len(questions))
    return questions


def parse_questions_from_docx(
    file,
    config: DocumentFormatConfig = None,
) -> list[dict[str, Any]]:
    """
    Parses questions from a DOCX file with configurable format
    """
    if config is None:
        config = DocumentFormatConfig()

    try:
        doc = Document(file)
        logger.info("DOCX file successfully read.")
    except Exception as e:
        logger.exception("Failed to read DOCX file: %s", str(e))
        raise

    # Extract all text content
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    # Check if it's the standard MCQ format (Question X) ... Ans.)
    if "Question" in content and "Ans." in content and "A)" in content:
        return _parse_mcq_format(content)

    # Standard configurable format parsing
    questions = []
    current_question = None
    current_options = []
    current_answer = None

    for para_num, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue

        # Check if it's a question
        if text.startswith(config.docx_question_prefix):
            # Save previous question if exists
            if current_question and current_options:
                question_data = _build_question_from_text(
                    current_question,
                    current_options,
                    current_answer,
                    config,
                )
                if question_data:
                    questions.append(question_data)

            # Start new question
            current_question = text[len(config.docx_question_prefix) :].strip()
            current_options = []
            current_answer = None

        # Check if it's an option
        elif any(text.startswith(prefix) for prefix in config.docx_option_prefixes):
            for i, prefix in enumerate(config.docx_option_prefixes):
                if text.startswith(prefix):
                    option_text = text[len(prefix) :].strip()
                    current_options.append(
                        {"index": i, "text": option_text, "letter": chr(ord("A") + i)},
                    )
                    break

        # Check if it's an answer
        elif text.startswith(config.docx_answer_prefix):
            current_answer = text[len(config.docx_answer_prefix) :].strip()

    # Don't forget the last question
    if current_question and current_options:
        question_data = _build_question_from_text(
            current_question,
            current_options,
            current_answer,
            config,
        )
        if question_data:
            questions.append(question_data)

    logger.info("Total valid questions parsed: %d", len(questions))
    return questions


def _parse_mcq_format(content: str) -> list[dict[str, Any]]:
    """
    Parse your standard MCQ format:
    Question 1) Which of the following is the primary purpose of a centrifuge in a laboratory setting?
    A) To measure the pH of a solution
    B) To separate components of a mixture based on density
    C) To sterilize laboratory equipment
    D) To amplify DNA sequences
    Ans. To separate components of a mixture based on density
    """
    questions = []

    # Split content by Question pattern
    question_pattern = r"Question\s+(\d+)\)"
    question_matches = list(re.finditer(question_pattern, content))

    for i, match in enumerate(question_matches):
        try:
            question_num = match.group(1)
            start_pos = match.end()

            # Find the end of this question (start of next question or end of content)
            if i + 1 < len(question_matches):
                end_pos = question_matches[i + 1].start()
            else:
                end_pos = len(content)

            question_block = content[start_pos:end_pos].strip()

            # Find the answer section
            if "Ans." not in question_block:
                logger.warning(f"Question {question_num}: No answer found, skipping")
                continue

            # Split question part and answer part
            parts = question_block.split("Ans.")
            if len(parts) != 2:
                logger.warning(f"Question {question_num}: Invalid format, skipping")
                continue

            question_part = parts[0].strip()
            answer_part = parts[1].strip()

            # Parse question text and options
            question_text = ""
            options = []

            # Split by newlines and process each line
            lines = question_part.split("\n")
            current_line = ""

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Check if it's an option (starts with A), B), C), D))
                option_match = re.match(r"^([A-D])\)\s*(.+)", line)
                if option_match:
                    option_letter = option_match.group(1)
                    option_text = option_match.group(2).strip()
                    options.append({"letter": option_letter, "text": option_text})
                # Part of question text
                elif current_line:
                    current_line += " " + line
                else:
                    current_line = line

            # The remaining current_line is the question text
            question_text = current_line.strip()

            # Find correct answer by matching with option text
            answer_text = answer_part.strip()
            correct_option = None

            # Try exact match first
            for option in options:
                if option["text"].strip() == answer_text:
                    correct_option = option
                    break

            # If no exact match, try case-insensitive match
            if not correct_option:
                for option in options:
                    if option["text"].lower().strip() == answer_text.lower().strip():
                        correct_option = option
                        break

            # If still no match, try partial match
            if not correct_option:
                for option in options:
                    if (
                        answer_text.lower().strip() in option["text"].lower()
                        or option["text"].lower() in answer_text.lower().strip()
                    ):
                        correct_option = option
                        break

            if not correct_option:
                logger.warning(
                    f"Question {question_num}: Could not match answer '{answer_text}' with any option",
                )
                continue

            # Build final question structure
            if question_text and options:
                answers = []
                for option in options:
                    answers.append(
                        {
                            "text": option["text"],
                            "option_letter": option["letter"],
                            "is_correct": (option == correct_option),
                        },
                    )

                questions.append(
                    {
                        "question": question_text,
                        "answers": answers,
                    },
                )

                logger.debug(
                    f"Question {question_num}: Parsed successfully with {len(answers)} options",
                )
            else:
                logger.warning(
                    f"Question {question_num}: Missing question text or options",
                )

        except Exception as e:
            logger.error(
                f"Question {question_num if 'question_num' in locals() else i + 1}: Error parsing - {e!s}",
            )
            continue

    logger.info(f"Total valid questions parsed from MCQ format: {len(questions)}")
    return questions


def _build_question_from_text(question_text, options, answer_text, config):
    """Helper function to build question data from text parsing"""
    if not question_text or not options:
        return None

    # Find correct answer
    correct_option = None
    if answer_text:
        if config.text_answer_format == "text":
            # Match by text content
            for option in options:
                if option["text"].lower() == answer_text.lower():
                    correct_option = option
                    break
        else:  # number format
            try:
                answer_num = int(answer_text) - 1  # Convert to 0-based index
                if 0 <= answer_num < len(options):
                    correct_option = options[answer_num]
            except ValueError:
                logger.warning("Invalid answer number: %s", answer_text)

    # Build answers list
    answers = []
    for option in options:
        answers.append(
            {
                "text": option["text"],
                "option_letter": option["letter"],
                "is_correct": (option == correct_option),
            },
        )

    return {
        "question": question_text,
        "answers": answers,
    }


def parse_questions_from_document(
    file,
    file_extension: str,
    config: DocumentFormatConfig = None,
) -> list[dict[str, Any]]:
    """
    Main function to parse questions from various document formats
    """
    file_extension = file_extension.lower()

    if file_extension == ".csv":
        return parse_questions_from_csv(file, config)
    if file_extension == ".txt":
        return parse_questions_from_text(file, config)
    if file_extension == ".docx":
        return parse_questions_from_docx(file, config)
    raise ValueError(f"Unsupported file format: {file_extension}")
